from docx import Document


def add_para(document, text):
    if (text.find("(BULLET_START)") == -1):
        document.add_paragraph(text)
    else:
        bullet_points = text.split("(BULLET_START)")
        for bullet_point in bullet_points:
            if (bullet_point.startswith(">>")):
                document.add_paragraph(bullet_point[2:], style='List Bullet')
            else:
                document.add_paragraph(bullet_point, style='List Bullet')

def execute():
    document = Document()
    with open("writable_file.txt", "+r", encoding="utf-8") as f:
        content = f.read()
        pages = content.split("(PAGE_BREAK)")

        for page in pages:
            try:
                if (page.find("(SECTION_BREAK)") == -1):
                    add_para(document, page)
                else:
                    sections = page.split("(SECTION_BREAK)")
                    for section in sections:
                        if (len(section.strip().strip("\n").strip("\r\n")) == 0):
                            continue
                        add_para(document, "\n\n" + section)
                        

                document.add_page_break()
            except:
                print(page)
                continue

    document.save("document.docx")